# pdf_parser.py
# Text extraction and preprocessing (Class 1-2 concepts)
#
# This is the first step in our NLP pipeline - before we can do any
# analysis (TF-IDF, extraction, etc.), we need clean text.
# Handles PDF, DOCX, and plain text files.

import fitz  # PyMuPDF
import re
from pathlib import Path


def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        raw = _extract_pdf(file_path)
    elif ext == ".docx":
        raw = _extract_docx(file_path)
    elif ext in (".txt", ".text"):
        raw = _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return clean_text(raw)


extract_text_from_pdf = extract_text  # legacy alias


def _extract_pdf(path: str) -> str:
    doc = fitz.open(path)
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    doc.close()
    return "\n\n".join(pages)


def _extract_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables (grading breakdowns, schedules often live in tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("  ".join(cells))

    return "\n\n".join(parts)


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def clean_text(text: str) -> str:
    # Text preprocessing / normalization (Class 1-2)
    # - normalize unicode chars (smart quotes, em-dashes, etc.)
    # - collapse whitespace but keep paragraph breaks
    # - remove page headers/footers
    # This is important because raw PDF text has a lot of encoding
    # artifacts that would mess up our regex patterns and TF-IDF
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"Syllabus for [A-Z]+-\d+, Page \d+\s*", "", text)
    return text.strip()
